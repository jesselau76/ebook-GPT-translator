from __future__ import annotations

import html
import os
import re
from pathlib import Path
from uuid import uuid4

import fitz
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import ITEM_DOCUMENT, ITEM_IMAGE, epub

from ebook_gpt_translator.config import AppConfig
from ebook_gpt_translator.models import Asset, Block, Chapter, Document


TEXTUAL_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote", "pre"}
_SENTENCE_BREAK_RE = re.compile(r"([。！？!?；;：:][\"'”’」』）】]*)")


def load_document(path: Path, config: AppConfig) -> Document:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _load_text_document(path)
    if suffix == ".docx":
        return _load_docx_document(path)
    if suffix == ".pdf":
        return _load_pdf_document(path, config)
    if suffix == ".epub":
        return _load_epub_document(path)
    if suffix == ".mobi":
        return _load_mobi_document(path)
    raise ValueError(f"Unsupported input format: {suffix}")


def _load_text_document(path: Path) -> Document:
    text = path.read_text(encoding="utf-8")
    blocks = []
    for index, paragraph in enumerate(_split_paragraphs(text), start=1):
        role = "heading" if index == 1 else "paragraph"
        blocks.append(Block(block_id=f"block-{index}", kind="text", role=role, text=paragraph))
    title = path.stem.replace("_", " ")
    return Document(source_path=path, format_name="txt", title=title, chapters=[Chapter("chapter-1", title, blocks=blocks)])


def _load_docx_document(path: Path) -> Document:
    doc = DocxDocument(path)
    blocks = []
    title = path.stem.replace("_", " ")
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        role = "heading" if style_name.startswith("heading") else "paragraph"
        blocks.append(Block(block_id=f"block-{index}", kind="text", role=role, text=text))
        if role == "heading" and title == path.stem.replace("_", " "):
            title = text
    return Document(source_path=path, format_name="docx", title=title, chapters=[Chapter("chapter-1", title, blocks=blocks)])


def _load_pdf_document(path: Path, config: AppConfig) -> Document:
    pdf = fitz.open(path)
    start_page = max(1, config.input.start_page)
    end_page = pdf.page_count if config.input.end_page in (-1, 0) else min(config.input.end_page, pdf.page_count)
    chapters: list[Chapter] = []
    for page_number in range(start_page - 1, end_page):
        page = pdf.load_page(page_number)
        page_text = page.get_text("text").strip()
        blocks = [
            Block(block_id=f"page-{page_number + 1}-block-{index}", kind="text", role="paragraph", text=paragraph)
            for index, paragraph in enumerate(_split_paragraphs(page_text), start=1)
        ]
        chapters.append(Chapter(chapter_id=f"page-{page_number + 1}", title=f"Page {page_number + 1}", blocks=blocks))
    pdf.close()
    return Document(source_path=path, format_name="pdf", title=path.stem.replace("_", " "), chapters=chapters)


def _load_epub_document(path: Path) -> Document:
    book = epub.read_epub(str(path))
    title = _first_or_default(book.get_metadata("DC", "title"), path.stem.replace("_", " "))
    author = _first_or_default(book.get_metadata("DC", "creator"), "")
    language = _first_or_default(book.get_metadata("DC", "language"), "en")

    assets: dict[str, Asset] = {}
    for item in book.get_items():
        if item.get_type() == ITEM_IMAGE:
            assets[item.get_name()] = Asset(
                asset_id=item.get_name(),
                href=item.get_name(),
                media_type=item.media_type,
                data=item.content,
            )

    chapters: list[Chapter] = []
    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue
        soup = BeautifulSoup(item.get_content(), "html.parser")
        body = soup.body or soup
        blocks: list[Block] = []
        for index, tag in enumerate(body.find_all(list(TEXTUAL_TAGS) + ["img"]), start=1):
            if tag.name == "img":
                src = (tag.get("src") or "").split("#", 1)[0]
                resolved = _resolve_asset_href(item.get_name(), src)
                if resolved in assets:
                    blocks.append(
                        Block(
                            block_id=f"{item.get_name()}-img-{index}",
                            kind="image",
                            role="image",
                            asset_id=resolved,
                        )
                    )
                continue
            text = tag.get_text(" ", strip=True)
            if not text:
                continue
            role = "heading" if tag.name.startswith("h") else "paragraph"
            blocks.append(Block(block_id=f"{item.get_name()}-text-{index}", kind="text", role=role, text=text))
        if blocks:
            chapter_title = next((block.text for block in blocks if block.role == "heading"), item.get_name())
            chapters.append(
                Chapter(
                    chapter_id=item.get_name().replace("/", "-"),
                    title=chapter_title,
                    source_href=item.get_name(),
                    blocks=blocks,
                )
            )
    return Document(source_path=path, format_name="epub", title=title, author=author, language=language, chapters=chapters, assets=assets)


def _load_mobi_document(path: Path) -> Document:
    try:
        import mobi  # type: ignore
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError("MOBI support requires the optional dependency: pip install mobi") from exc

    tmp_dir, _ = mobi.extract(str(path))
    html_path = next(Path(tmp_dir).rglob("*.html"), None)
    if html_path is None:
        html_path = next(Path(tmp_dir).rglob("*.xhtml"), None)
    if html_path is None:
        raise RuntimeError("Unable to extract HTML content from MOBI file")
    html_content = html_path.read_text(encoding="utf-8", errors="ignore")
    temp_source = path.with_suffix(".html")
    temp_source.write_text(html_content, encoding="utf-8")
    try:
        return _load_html_document(temp_source, path)
    finally:
        if temp_source.exists():
            temp_source.unlink()


def _load_html_document(temp_path: Path, original_path: Path | None = None) -> Document:
    source_path = original_path or temp_path
    soup = BeautifulSoup(temp_path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    body = soup.body or soup
    blocks = []
    for index, tag in enumerate(body.find_all(list(TEXTUAL_TAGS)), start=1):
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        role = "heading" if tag.name.startswith("h") else "paragraph"
        blocks.append(Block(block_id=f"block-{index}", kind="text", role=role, text=text))
    title = soup.title.get_text(strip=True) if soup.title else source_path.stem.replace("_", " ")
    return Document(source_path=source_path, format_name=source_path.suffix.lower().lstrip("."), title=title, chapters=[Chapter("chapter-1", title, blocks=blocks)])


def write_outputs(document: Document, config: AppConfig) -> tuple[Path | None, Path | None]:
    output_dir = Path(config.output.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    txt_path: Path | None = None
    epub_path: Path | None = None
    stem = document.source_path.stem
    if config.output.emit_txt:
        txt_path = output_dir / f"{stem}.translated.txt"
        _ensure_writable(txt_path, config.output.overwrite)
        txt_path.write_text(_render_txt(document, config.translation.bilingual_output), encoding="utf-8")
    if config.output.emit_epub:
        epub_path = output_dir / f"{stem}.translated.epub"
        _ensure_writable(epub_path, config.output.overwrite)
        _write_epub(document, epub_path, config.translation.bilingual_output)
    return txt_path, epub_path


def _render_txt(document: Document, bilingual_output: bool) -> str:
    parts: list[str] = []
    for chapter in document.chapters:
        title = chapter.translated_title or chapter.title
        if title:
            parts.append(title)
            parts.append("=" * len(title))
        for block in chapter.blocks:
            if block.kind == "image":
                parts.append(f"[Image: {block.asset_id}]")
                continue
            translated = _format_reading_text(block.translated_text or block.text)
            if bilingual_output:
                parts.append(block.text)
                parts.append(translated)
            else:
                parts.append(translated)
            parts.append("")
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def _write_epub(document: Document, output_path: Path, bilingual_output: bool) -> None:
    book = epub.EpubBook()
    book.set_identifier(str(uuid4()))
    book.set_title(document.title)
    book.set_language(document.language or "en")
    if document.author:
        book.add_author(document.author)

    style = epub.EpubItem(
        uid="style-nav",
        file_name="style/style.css",
        media_type="text/css",
        content="""
        body { font-family: serif; line-height: 1.5; }
        .source { color: #666; font-style: italic; }
        .translated { margin-bottom: 1em; }
        img { max-width: 100%; }
        """.encode("utf-8"),
    )
    book.add_item(style)

    for asset in document.assets.values():
        book.add_item(epub.EpubItem(uid=asset.asset_id, file_name=asset.href, media_type=asset.media_type, content=asset.data))

    spine = ["nav"]
    toc = []
    for chapter_index, chapter in enumerate(document.chapters, start=1):
        title = chapter.translated_title or chapter.title or f"Chapter {chapter_index}"
        html_parts = [f"<h1>{html.escape(title)}</h1>"]
        for block in chapter.blocks:
            if block.kind == "image":
                if block.asset_id:
                    html_parts.append(f'<p><img src="{html.escape(document.assets[block.asset_id].href)}" alt="" /></p>')
                continue
            translated = _format_epub_text(block.translated_text or block.text)
            if bilingual_output:
                source = html.escape(block.text)
                html_parts.append(f'<p class="source">{source}</p>')
                html_parts.append(f'<p class="translated">{translated}</p>')
            elif block.heading_text:
                # Merged block with heading: split first paragraph as <h2>, rest as <p>
                parts = translated.split("<br /><br />", 1)
                html_parts.append(f"<h2>{parts[0]}</h2>")
                if len(parts) > 1:
                    html_parts.append(f"<p>{parts[1]}</p>")
            else:
                tag = "p" if block.role == "paragraph" else "h2"
                html_parts.append(f"<{tag}>{translated}</{tag}>")
        epub_chapter = epub.EpubHtml(
            title=title,
            file_name=chapter.source_href or f"chapter-{chapter_index}.xhtml",
            lang=document.language or "en",
        )
        epub_chapter.content = (
            "<html><head><link rel='stylesheet' type='text/css' href='style/style.css' /></head>"
            f"<body>{''.join(html_parts)}</body></html>"
        )
        book.add_item(epub_chapter)
        toc.append(epub_chapter)
        spine.append(epub_chapter)

    book.toc = tuple(toc)
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(output_path), book)


def _split_paragraphs(text: str) -> list[str]:
    return [paragraph.strip() for paragraph in text.splitlines() if paragraph.strip()]


def _format_reading_text(text: str) -> str:
    normalized = "\n".join(line.rstrip() for line in text.strip().splitlines())
    if not normalized:
        return ""
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = _SENTENCE_BREAK_RE.sub(r"\1\n\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _format_epub_text(text: str) -> str:
    return html.escape(_format_reading_text(text)).replace("\n\n", "<br /><br />").replace("\n", "<br />")


def _resolve_asset_href(base_href: str, src: str) -> str:
    normalized = os.path.normpath(os.path.join(os.path.dirname(base_href), src)).replace("\\", "/")
    return normalized


def _first_or_default(values: list[tuple[str, dict]] | list, default: str) -> str:
    if not values:
        return default
    first = values[0]
    if isinstance(first, tuple):
        return first[0]
    return str(first)


def _ensure_writable(path: Path, overwrite: bool) -> None:
    if path.exists() and not overwrite:
        raise FileExistsError(
            f"Output already exists: {path}. Use --overwrite or --skip-existing to control this behavior."
        )
