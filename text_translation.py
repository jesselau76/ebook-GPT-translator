# -*- coding: utf-8 -*-

import pdfminer.high_level
import re
import openai
from tqdm import tqdm
# import nltk
# nltk.download('punkt')
# from nltk.tokenize import sent_tokenize
import ebooklib
from ebooklib import epub
import os
import tempfile
import shutil
from bs4 import BeautifulSoup
import configparser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
import random
import json
import docx
import zipfile
from lxml import etree
from docx import Document
import mobi
import pandas as pd


def get_docx_title(docx_filename):
    with zipfile.ZipFile(docx_filename) as zf:
        core_properties = etree.fromstring(zf.read("docProps/core.xml"))

    ns = {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
          "dc": "http://purl.org/dc/elements/1.1/",
          "dcterms": "http://purl.org/dc/terms/",
          "dcmitype": "http://purl.org/dc/dcmitype/",
          "xsi": "http://www.w3.org/2001/XMLSchema-instance"}

    title_elements = core_properties.findall("dc:title", ns)
    if title_elements:
        return title_elements[0].text
    else:
        return "Unknown title"


def get_pdf_title(pdf_filename):
    try:
        with open(pdf_filename, 'rb') as file:
            parser = PDFParser(file)
            document = PDFDocument(parser)
            if 'Title' in document.info:
                return document.info['Title']
            else:
                text = pdfminer.high_level.extract_text(file)
                match = re.search(r'(?<=\n)([^\n]+)(?=\n)', text)
                if match:
                    return match.group(1)
                else:
                    return "Unknown title"
    except:
        return "Unknown title"


def get_mobi_title(mobi_filename):
    try:
        metadata = mobi.read_metadata(mobi_filename)
        title = metadata.get("Title", None)
    except:
        return "Unknown title"


def convert_mobi_to_text(mobi_filename):
    # Extract MOBI contents to a temporary directory
    with tempfile.TemporaryDirectory() as tempdir:
        tempdir, filepath = mobi.extract(mobi_filename)

        # Find the HTML file in the temporary directory
        for root, _, files in os.walk(tempdir):
            for file in files:
                if file.endswith(".html"):
                    html_file = os.path.join(root, file)
                    break
            else:
                continue
            break
        else:
            raise FileNotFoundError("HTML file not found in the extracted MOBI contents")

        # Parse the HTML file with BeautifulSoup to get the text
        with open(html_file, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            text = soup.get_text()

    return text


def get_epub_title(epub_filename):
    try:
        book = epub.read_epub(epub_filename)
        metadata = book.get_metadata('DC', {})
        if metadata:
            if 'title' in metadata:
                return metadata['title'][0]
        else:
            return "Unknown title"
    except:
        return "Unknown title"

    # 读取option文件


import chardet

with open('settings.cfg', 'rb') as f:
    content = f.read()
    encoding = chardet.detect(content)['encoding']

with open('settings.cfg', encoding=encoding) as f:
    config_text = f.read()
    config = configparser.ConfigParser()
    config.read_string(config_text)

# 获取openai_apikey和language
openai_apikey = config.get('option', 'openai-apikey')
# language_name = config.get('option', 'target-language')
prompt = config.get('option', 'prompt')
bilingual_output = config.get('option', 'bilingual-output')
language_code = config.get('option', 'langcode')
api_proxy=config.get('option', 'openai-proxy')
# Get startpage and endpage as integers with default values
startpage = config.getint('option', 'startpage', fallback=1)
endpage = config.getint('option', 'endpage', fallback=-1)
# 设置译名表文件路径
transliteration_list_file = config.get('option', 'transliteration-list')
# 译名表替换是否开启大小写匹配？
case_matching = config.get('option', 'case-matching')

# 设置openai的API密钥
openai.api_key = openai_apikey

# 将openai的API密钥分割成数组
key_array = openai_apikey.split(',')

def random_api_key():
    return random.choice(key_array)

def create_chat_completion(prompt, text, model="gpt-3.5-turbo", **kwargs):
    openai.api_key = random_api_key()
    return openai.ChatCompletion.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": f"{prompt}: \n{text}",
            }
        ],
        **kwargs
    )

import argparse

# 如果配置文件有写，就设置api代理
if len(api_proxy) == 0:
    print("未检测到OpenAI API 代理，当前使用api地址为: " + openai.api_base)
else:
    api_proxy_url = api_proxy + "/v1"
    openai.api_base = os.environ.get("OPENAI_API_BASE", api_proxy_url)
    print("正在使用OpenAI API 代理，代理地址为: "+openai.api_base)

# 创建参数解析器
parser = argparse.ArgumentParser()
parser.add_argument("filename", help="Name of the input file")
parser.add_argument("--test", help="Only translate the first 3 short texts", action="store_true")
# 是否使用译名表？
parser.add_argument("--tlist", help="Use the translated name table", action="store_true")
args = parser.parse_args()

# 获取命令行参数
filename = args.filename
base_filename, file_extension = os.path.splitext(filename)
new_filename = base_filename + "_translated.epub"
new_filenametxt = base_filename + "_translated.txt"
jsonfile = base_filename + "_process.json"
# 从文件中加载已经翻译的文本
translated_dict = {}
try:
    with open(jsonfile, "r", encoding="utf-8") as f:
        translated_dict = json.load(f)
except FileNotFoundError:
    pass


def convert_docx_to_text(docx_filename):
    doc = docx.Document(docx_filename)

    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return text


def convert_epub_to_text(epub_filename):
    # 打开epub文件
    book = epub.read_epub(epub_filename)

    # 获取所有文本
    text = ""
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            # 使用BeautifulSoup提取纯文本
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text += re.sub(r'\n+', '\n', soup.get_text().strip())

    # 返回文本
    return text


def text_to_epub(text, filename, language_code='en', title="Title"):
    text = text.replace("\n", "<br>")
    # 创建epub书籍对象
    book = epub.EpubBook()

    # 设置元数据
    book.set_identifier(str(random.randint(100000, 999999)))
    book.set_title(title)
    book.set_language(language_code)

    # 创建章节对象
    c = epub.EpubHtml(title='Chapter 1', file_name='chap_1.xhtml', lang=language_code)
    c.content = text

    # 将章节添加到书籍中
    book.add_item(c)

    # 添加toc
    book.toc = (epub.Link('chap_1.xhtml', 'Chapter 1', 'chap_1'),)
    # 设置书脊顺序
    book.spine = ['nav', c]
    # 添加导航
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # 设置书籍封面
    # book.set_cover('image.jpg', open('image.jpg', 'rb').read())

    # 将书籍写入文件
    epub.write_epub(filename, book, {})


# 将PDF文件转换为文本
# For PDF files
def get_total_pages(pdf_filename):
    with open(pdf_filename, 'rb') as file:
        parser = PDFParser(file)
        document = PDFDocument(parser)
        return len(list(PDFPage.create_pages(document)))


def convert_pdf_to_text(pdf_filename, start_page=1, end_page=-1):
    if end_page == -1:
        end_page = get_total_pages(pdf_filename)
        # print("Total pages of the file:"+ str(end_page))
        # print("Converting PDF from:"+ str(start_page)+" to "+ str(end_page) + " page")
        text = pdfminer.high_level.extract_text(pdf_filename, page_numbers=list(range(start_page - 1, end_page)))
    else:
        # print("Converting PDF from:"+ str(start_page)+" to "+ str(end_page) + " page")
        text = pdfminer.high_level.extract_text(pdf_filename, page_numbers=list(range(start_page - 1, end_page)))
    return text


# 将文本分成不大于1024字符的短文本list
def split_text(text):
    sentence_list = re.findall(r'.+?[。！？!?.]', text)

    # 初始化短文本列表
    short_text_list = []
    # 初始化当前短文本
    short_text = ""
    # 遍历句子列表
    for s in sentence_list:
        # 如果当前短文本加上新的句子长度不大于1024，则将新的句子加入当前短文本
        if len(short_text + s) <= 1024:
            short_text += s
        # 如果当前短文本加上新的句子长度大于1024，则将当前短文本加入短文本列表，并重置当前短文本为新的句子
        else:
            short_text_list.append(short_text)
            short_text = s
    # 将最后的短文本加入短文本列表
    short_text_list.append(short_text)
    return short_text_list


# 将句号替换为句号+回车
def return_text(text):
    text = text.replace(". ", ".\n")
    text = text.replace("。", "。\n")
    text = text.replace("！", "！\n")
    return text


# Initialize a count variable of tokens cost.
cost_tokens = 0


# 翻译短文本
def translate_text(text):
    global cost_tokens

    # 调用openai的API进行翻译
    try:
        completion = create_chat_completion(prompt, text)
        t_text = (
            completion["choices"][0]
            .get("message")
            .get("content")
            .encode("utf8")
            .decode()
        )
        # Get the token usage from the API response
        cost_tokens += completion["usage"]["total_tokens"]

    except Exception as e:
        import time
        # TIME LIMIT for open api please pay
        sleep_time = 60
        time.sleep(sleep_time)
        print(e, f"will sleep  {sleep_time} seconds")

        completion = create_chat_completion(prompt, text)
        t_text = (
            completion["choices"][0]
            .get("message")
            .get("content")
            .encode("utf8")
            .decode()
        )
        # Get the token usage from the API response
        cost_tokens += completion["usage"]["total_tokens"]

    return t_text


def translate_and_store(text):
    # 如果文本已经翻译过，直接返回翻译结果
    if text in translated_dict:
        return translated_dict[text]

    # 否则，调用 translate_text 函数进行翻译，并将结果存储在字典中
    translated_text = translate_text(text)
    translated_dict[text] = translated_text

    # 将字典保存为 JSON 文件
    with open(jsonfile, "w", encoding="utf-8") as f:
        json.dump(translated_dict, f, ensure_ascii=False, indent=4)

    return translated_text


def text_replace(long_string, xlsx_path, case_sensitive):
    # 读取excel文件，将第一列和第二列分别存为两个列表
    df = pd.read_excel(xlsx_path)
    old_words = df.iloc[:, 0].tolist()
    new_words = df.iloc[:, 1].tolist()
    # 对旧词列表按照长度降序排序，并同步调整新词列表
    old_words, new_words = zip(*sorted(zip(old_words, new_words), key=lambda x: len(x[0]), reverse=True))
    # 遍历两个列表，对字符串进行替换
    for i in range(len(old_words)):
        # 如果不区分大小写，就将字符串和被替换词都转为小写
        if not case_sensitive:
            lower_string = long_string.lower()
            lower_old_word = old_words[i].lower()
            # 使用正则表达式进行替换，注意要保留原字符串的大小写
            long_string = re.sub(r"\b" + lower_old_word + r"\b", new_words[i], long_string, flags=re.IGNORECASE)
        # 如果区分大小写，就直接使用正则表达式进行替换
        else:
            long_string = re.sub(r"\b" + old_words[i] + r"\b", new_words[i], long_string)
    # 返回替换后的字符串
    return long_string


text = ""

# 根据文件类型调用相应的函数
if filename.endswith('.pdf'):
    print("Converting PDF to text")
    title = get_pdf_title(filename)
    with tqdm(total=10, desc="Converting PDF to text") as pbar:
        for i in range(10):
            text = convert_pdf_to_text(filename, startpage, endpage)
            pbar.update(1)
elif filename.endswith('.epub'):
    print("Converting epub to text")
    book = epub.read_epub(filename)
elif filename.endswith('.txt'):

    with open(filename, 'r', encoding='utf-8') as file:
        text = file.read()

    title = os.path.basename(filename)
elif filename.endswith('.docx'):
    print("Converting DOCX file to text")
    title = get_docx_title(filename)
    with tqdm(total=10, desc="Converting DOCX to text") as pbar:
        for i in range(10):
            text = convert_docx_to_text(filename)
            pbar.update(1)

elif filename.endswith('.mobi'):
    print("Converting MOBI file to text")
    title = get_mobi_title(filename)
    with tqdm(total=10, desc="Converting MOBI to text") as pbar:
        for i in range(10):
            text = convert_mobi_to_text(filename)
            pbar.update(1)
else:
    print("Unsupported file type")

if filename.endswith('.epub'):
    # 获取所有章节
    items = book.get_items()

    # 遍历所有章节
    translated_all = ''
    count = 0
    for item in tqdm(items):
        # 如果章节类型为文档类型，则需要翻译
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            # 使用BeautifulSoup提取原文本
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text().strip()
            img_html = ''
            img_tags = soup.find_all('img')
            for img_tag in img_tags:
                img_html += str(img_tag) + '<br>'
            # 如果原文本为空，则跳过
            if not text:
                continue
            # 将所有回车替换为空格
            text = text.replace("\n", " ")

            # 将多个空格替换为一个空格
            import re

            text = re.sub(r"\s+", " ", text)

            # 如果设置了译名表替换，则对文本进行翻译前的替换
            if args.tlist:
                text = text_replace(text, transliteration_list_file, case_matching)

            # 将文本分成不大于1024字符的短文本list
            short_text_list = split_text(text)
            if args.test:
                short_text_list = short_text_list[:3]

            # 初始化翻译后的文本
            translated_text = ""

            # 遍历短文本列表，依次翻译每个短文本
            for short_text in tqdm(short_text_list):
                print(return_text(short_text))
                count += 1
                # 翻译当前短文本
                translated_short_text = translate_and_store(short_text)
                short_text = return_text(short_text)
                translated_short_text = return_text(translated_short_text)
                # 将当前短文本和翻译后的文本加入总文本中
                if bilingual_output.lower() == 'true':
                    translated_text += f"{short_text}<br>\n{translated_short_text}<br>\n"
                else:
                    translated_text += f"{translated_short_text}<br>\n"
                # print(short_text)
                print(translated_short_text)
            # 使用翻译后的文本替换原有的章节内容
            item.set_content((img_html + translated_text.replace('\n', '<br>')).encode('utf-8'))
            translated_all += translated_text
            if args.test and count >= 3:
                break

    # 将epub书籍写入文件
    epub.write_epub(new_filename, book, {})
    # 将翻译后的文本同时写入txt文件 in case epub插件出问题
    with open(new_filenametxt, "w", encoding="utf-8") as f:
        f.write(translated_all)

else:
    # 将所有回车替换为空格
    text = text.replace("\n", " ")

    # 将多个空格替换为一个空格
    import re

    text = re.sub(r"\s+", " ", text)

    # 如果设置了译名表替换，则对文本进行翻译前的替换
    if args.tlist:
        text = text_replace(text, transliteration_list_file, case_matching)

    # 将文本分成不大于1024字符的短文本list
    short_text_list = split_text(text)
    if args.test:
        short_text_list = short_text_list[:3]
    # 初始化翻译后的文本
    translated_text = ""

    # 遍历短文本列表，依次翻译每个短文本
    for short_text in tqdm(short_text_list):
        print(return_text(short_text))
        # 翻译当前短文本
        translated_short_text = translate_and_store(short_text)
        short_text = return_text(short_text)
        translated_short_text = return_text(translated_short_text)
        # 将当前短文本和翻译后的文本加入总文本中
        if bilingual_output.lower() == 'true':
            translated_text += f"{short_text}\n{translated_short_text}\n"
        else:
            translated_text += f"{translated_short_text}\n"
        # print(short_text)
        print(translated_short_text)

    # 将翻译后的文本写入epub文件
    with tqdm(total=10, desc="Writing translated text to epub") as pbar:
        text_to_epub(translated_text.replace('\n', '<br>'), new_filename, language_code, title)
        pbar.update(1)

    # 将翻译后的文本同时写入txt文件 in case epub插件出问题
    with open(new_filenametxt, "w", encoding="utf-8") as f:
        f.write(translated_text)
cost = cost_tokens / 1000 * 0.002
print(f"Translation completed. Total cost: {cost_tokens} tokens, ${cost}.")

try:
    os.remove(jsonfile)
    print(f"File '{jsonfile}' has been deleted.")
except FileNotFoundError:
    print(f"File '{jsonfile}' not found. No file was deleted.")
